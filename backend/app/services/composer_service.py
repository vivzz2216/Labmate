import os
import uuid
import re
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
from sqlalchemy.orm import Session
from ..models import Upload, AITask, AIJob
from ..config import settings


class ComposerService:
    """Service for composing final DOCX reports with embedded screenshots"""
    
    def __init__(self):
        pass
    
    async def compose_report(
        self, 
        upload_id: int, 
        screenshot_order: List[int], 
        db: Session
    ) -> Dict[str, Any]:
        """
        Update the original document with embedded screenshots
        
        Args:
            upload_id: ID of the uploaded document
            screenshot_order: Ordered list of task IDs for screenshots
            db: Database session
            
        Returns:
            Dict with report_path, filename, and download_url
        """
        
        # Get the original upload
        upload = db.query(Upload).filter(Upload.id == upload_id).first()
        if not upload:
            raise ValueError("Upload not found")
        
        # Get tasks with screenshots in the specified order
        tasks_with_screenshots = []
        for task_id in screenshot_order:
            task = db.query(AITask).filter(AITask.id == task_id).first()
            if task and task.screenshot_path and task.status == 'completed':
                tasks_with_screenshots.append(task)
        
        # Load the original document
        if not os.path.exists(upload.file_path):
            raise ValueError("Original document not found")
        
        # Create a copy of the original document
        doc = Document(upload.file_path)
        
        # Add a header indicating screenshots have been added
        header_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        
        # Insert screenshots at the end of the document
        await self._add_screenshots_to_document(doc, tasks_with_screenshots)
        
        # Save the updated document
        report_filename = f"{upload.filename.replace('.docx', '')}_with_screenshots_{uuid.uuid4().hex[:8]}.docx"
        report_path = os.path.join(settings.REPORT_DIR, report_filename)
        
        doc.save(report_path)
        
        return {
            "report_path": report_path,
            "filename": report_filename,
            "download_url": f"/reports/{report_filename}"
        }
    
    async def _add_screenshots_to_document(self, doc: Document, tasks_with_screenshots: List[AITask]):
        """Add screenshots to the original document - try to place near questions"""
        
        # Try to insert screenshots near relevant questions
        await self._insert_screenshots_near_questions(doc, tasks_with_screenshots)
    
    async def _insert_screenshots_near_questions(self, doc: Document, tasks_with_screenshots: List[AITask]):
        """Insert screenshots under relevant questions in the document"""
        
        if not tasks_with_screenshots:
            return
        
        # Create a mapping of task numbers to tasks
        task_mapping = {}
        for task in tasks_with_screenshots:
            # Extract task number from brief_description or use sequential numbering
            task_num = self._extract_task_number(task.brief_description) if task.brief_description else len(task_mapping) + 1
            task_mapping[task_num] = task
        
        # Find all paragraphs and try to insert screenshots after matching questions
        paragraphs_to_insert = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Look for question patterns
            task_num = self._find_question_pattern(text)
            if task_num and task_num in task_mapping:
                task = task_mapping[task_num]
                # Find the end of this question (look for next numbered question or section)
                end_index = self._find_question_end_index(doc.paragraphs, i)
                # Mark this paragraph for screenshot insertion
                paragraphs_to_insert.append((end_index, task))  # Insert after the question
                # Remove from mapping to avoid duplicate insertions
                del task_mapping[task_num]
        
        # Insert screenshots after the identified paragraphs (in reverse order to maintain indices)
        for para_index, task in reversed(paragraphs_to_insert):
            await self._insert_screenshot_after_paragraph(doc, para_index, task)
        
        # Add any remaining screenshots at the end (fallback) - question + screenshot pairs
        if task_mapping:
            doc.add_paragraph()
            doc.add_paragraph("Additional Solutions:")
            for task_num, task in task_mapping.items():
                # Add the question text
                if task.brief_description:
                    question_para = doc.add_paragraph()
                    question_para.add_run(f"Question {task_num}: ").bold = True
                    question_para.add_run(task.brief_description)
                
                # Add the screenshot
                if task.screenshot_path and os.path.exists(task.screenshot_path):
                    await self._add_screenshot_only_clean(doc, task.screenshot_path, f"Solution {task_num}")
                
                # Add small spacing between questions
                doc.add_paragraph()
    
    def _find_question_end_index(self, paragraphs, start_index: int) -> int:
        """Find where a question ends (next numbered item or section)"""
        for i in range(start_index + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            
            # Check if this is a new numbered question
            if self._find_question_pattern(text):
                return i
            
            # Check if this is a new section (like "C. Questions:", "A. Theory:", etc.)
            if re.match(r'^[A-Z]\.\s+[A-Z]', text):
                return i
        
        # If no end found, return the next paragraph index
        return start_index + 1
    
    def _extract_task_number(self, description: str) -> Optional[int]:
        """Extract task number from description"""
        if not description:
            return None
        
        # Look for patterns like "Task 1", "Question 1", "1.", etc.
        patterns = [
            r'Task\s+(\d+)',
            r'Question\s+(\d+)',
            r'Problem\s+(\d+)',
            r'Exercise\s+(\d+)',
            r'^(\d+)\.',
            r'^(\d+)\)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        return None
    
    def _find_question_pattern(self, text: str) -> Optional[int]:
        """Find question pattern in text and return task number"""
        patterns = [
            r'Question\s+(\d+)',
            r'Task\s+(\d+)',
            r'Problem\s+(\d+)',
            r'Exercise\s+(\d+)',
            r'^(\d+)\.',                    # "1.Write a Python program..."
            r'^(\d+)\)',                    # "1) Write a Python program..."
            r'Q(\d+)',                      # "Q1", "Q2"
            r'T(\d+)',                      # "T1", "T2"
            r'^\s*(\d+)\.\s*Write',         # "1. Write a Python program..."
            r'^\s*(\d+)\)\s*Write',         # "1) Write a Python program..."
            r'^\s*(\d+)\.\s*[A-Z]',         # "1. Demonstrate", "2. Calculate"
            r'^\s*(\d+)\)\s*[A-Z]'          # "1) Demonstrate", "2) Calculate"
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return int(match.group(1))
        
        return None
    
    async def _insert_screenshot_after_paragraph(self, doc: Document, para_index: int, task: AITask):
        """Insert screenshot after a specific paragraph - only the screenshot, no extra text"""
        if task.screenshot_path and os.path.exists(task.screenshot_path):
            # Add the screenshot with minimal caption
            task_num = self._extract_task_number(task.brief_description)
            caption = f"Task {task_num}" if task_num else "Code Execution"
            await self._add_screenshot_only_clean(doc, task.screenshot_path, caption)
    
    async def _add_screenshot_only(self, doc: Document, task: AITask, task_number: int):
        """Add just the screenshot with minimal text"""
        
        # Add screenshot
        if task.screenshot_path and os.path.exists(task.screenshot_path):
            await self._add_screenshot_only_clean(doc, task.screenshot_path, f"Task {task_number}")
    
    async def _add_screenshot_only_clean(self, doc: Document, image_path: str, caption: str):
        """Add only the screenshot with minimal caption - no extra text or formatting"""
        
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                width, height = img.size
            
            # Calculate size to fit within document (max 6 inches width)
            max_width = Inches(6)
            max_height = Inches(4)
            
            # Calculate aspect ratio
            aspect_ratio = width / height
            
            if width > max_width:
                new_width = max_width
                new_height = new_width / aspect_ratio
            else:
                new_width = Inches(width / 100)  # Convert pixels to inches (approximate)
                new_height = Inches(height / 100)
            
            if new_height > max_height:
                new_height = max_height
                new_width = new_height * aspect_ratio
            
            # Add image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image to paragraph
            doc.add_picture(image_path, width=new_width, height=new_height)
            
            # Add minimal caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            
        except Exception as e:
            # If image fails to load, add minimal error message
            error_para = doc.add_paragraph()
            error_para.add_run(f"Screenshot unavailable")
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def _add_task_section(self, doc: Document, task: AITask, task_number: int):
        """Add a task section with screenshot and details to the document"""
        
        # Task heading
        task_heading = doc.add_heading(f'Task {task_number}', level=2)
        
        # Task description
        if task.brief_description:
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(task.brief_description)
        
        # Add screenshot
        if task.screenshot_path and os.path.exists(task.screenshot_path):
            await self._add_screenshot(doc, task.screenshot_path, f"Task {task_number} Screenshot")
        
        # Add execution output
        if task.stdout:
            output_heading = doc.add_heading('Execution Output', level=3)
            output_para = doc.add_paragraph()
            
            # Create a code-like formatting
            output_para.style = 'Normal'
            output_run = output_para.add_run(task.stdout)
            output_run.font.name = 'Courier New'
            output_run.font.size = Pt(10)
            
            # Add background color (gray)
            output_para.paragraph_format.left_indent = Inches(0.5)
            output_para.paragraph_format.right_indent = Inches(0.5)
            output_para.paragraph_format.space_before = Pt(6)
            output_para.paragraph_format.space_after = Pt(6)
        
        # Add AI answer if available
        if task.assistant_answer:
            answer_heading = doc.add_heading('AI Explanation', level=3)
            answer_para = doc.add_paragraph(task.assistant_answer)
        
        # Add status
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f'Status: {task.status.upper()}')
        status_run.bold = True
        status_run.font.color.rgb = self._get_status_color(task.status)
        
        # Add separator between tasks
        doc.add_paragraph()
        doc.add_paragraph('-' * 30)
        doc.add_paragraph()
    
    async def _add_screenshot(self, doc: Document, image_path: str, caption: str):
        """Add a screenshot to the document with proper sizing and caption"""
        
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                width, height = img.size
            
            # Calculate size to fit within document (max 6 inches width)
            max_width = Inches(6)
            max_height = Inches(4)
            
            # Calculate aspect ratio
            aspect_ratio = width / height
            
            if width > max_width:
                new_width = max_width
                new_height = new_width / aspect_ratio
            else:
                new_width = Inches(width / 100)  # Convert pixels to inches (approximate)
                new_height = Inches(height / 100)
            
            if new_height > max_height:
                new_height = max_height
                new_width = new_height * aspect_ratio
            
            # Add image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # Add image to paragraph
            doc.add_picture(image_path, width=new_width, height=new_height)
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            
        except Exception as e:
            # If image fails to load, add error message
            error_para = doc.add_paragraph()
            error_para.add_run(f"Error loading screenshot: {str(e)}")
            error_para.runs[0].font.color.rgb = self._get_status_color('failed')
    
    def _get_status_color(self, status: str):
        """Get RGB color for status"""
        from docx.shared import RGBColor
        
        if status == 'completed':
            return RGBColor(0, 128, 0)  # Green
        elif status == 'failed':
            return RGBColor(255, 0, 0)  # Red
        else:
            return RGBColor(128, 128, 128)  # Gray


# Create singleton instance
composer_service = ComposerService()
