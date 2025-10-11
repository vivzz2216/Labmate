import os
import uuid
from datetime import datetime
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sqlalchemy.orm import Session
from ..models import Upload, Job, Screenshot
from ..config import settings


class DocxService:
    """Service for composing final DOCX reports with embedded screenshots"""
    
    def __init__(self):
        self.report_dir = settings.REPORT_DIR
    
    async def compose_report(
        self, 
        upload_id: int, 
        screenshot_order: Optional[List[int]] = None,
        db: Session = None
    ) -> Tuple[bool, str, str]:
        """
        Compose final DOCX report with embedded screenshots
        
        Args:
            upload_id: ID of the uploaded file
            screenshot_order: Ordered list of job IDs for screenshot placement
            db: Database session
            
        Returns:
            Tuple of (success, file_path, filename)
        """
        
        try:
            # Get upload record
            upload = db.query(Upload).filter(Upload.id == upload_id).first()
            if not upload:
                return False, "", "Upload not found"
            
            # Load original document
            doc = Document(upload.file_path)
            
            # Get jobs and screenshots
            jobs = db.query(Job).filter(Job.upload_id == upload_id).all()
            
            if screenshot_order:
                # Sort jobs by provided order
                job_dict = {job.id: job for job in jobs}
                ordered_jobs = [job_dict[job_id] for job_id in screenshot_order if job_id in job_dict]
            else:
                # Use default order (by task_id)
                ordered_jobs = sorted(jobs, key=lambda x: x.task_id)
            
            # Process each job and insert screenshots
            await self._insert_screenshots(doc, ordered_jobs, db)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_final_{timestamp}.docx"
            file_path = os.path.join(self.report_dir, filename)
            
            # Ensure report directory exists
            os.makedirs(self.report_dir, exist_ok=True)
            
            # Save document
            doc.save(file_path)
            
            return True, file_path, filename
            
        except Exception as e:
            return False, "", f"Report composition error: {str(e)}"
    
    async def _insert_screenshots(
        self, 
        doc: Document, 
        jobs: List[Job], 
        db: Session
    ):
        """Insert screenshots into the document at appropriate locations"""
        
        for job in jobs:
            # Get screenshot for this job
            screenshot = db.query(Screenshot).filter(Screenshot.job_id == job.id).first()
            if not screenshot or not os.path.exists(screenshot.file_path):
                continue
            
            # Find the best location to insert the screenshot
            insertion_point = await self._find_insertion_point(doc, job.question_text)
            
            if insertion_point is not None:
                # Insert screenshot at the found location
                await self._insert_screenshot_at_paragraph(doc, screenshot, insertion_point, job.task_id)
    
    async def _find_insertion_point(self, doc: Document, question_text: str) -> Optional[int]:
        """Find the best paragraph to insert screenshot after"""
        
        # Clean question text for matching
        clean_question = self._clean_text_for_matching(question_text)
        
        best_match_index = None
        best_match_score = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
            
            # Clean paragraph text
            clean_para = self._clean_text_for_matching(paragraph.text)
            
            # Calculate similarity score
            score = self._calculate_similarity(clean_question, clean_para)
            
            if score > best_match_score and score > 0.3:  # Minimum similarity threshold
                best_match_score = score
                best_match_index = i
        
        return best_match_index
    
    def _clean_text_for_matching(self, text: str) -> str:
        """Clean text for similarity matching"""
        import re
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Remove punctuation for better matching
        text = re.sub(r'[^\w\s]', '', text)
        
        # Convert to lowercase
        text = text.lower()
        
        return text
    
    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate simple text similarity score"""
        if not text1 or not text2:
            return 0.0
        
        # Split into words
        words1 = set(text1.split())
        words2 = set(text2.split())
        
        if not words1 or not words2:
            return 0.0
        
        # Calculate Jaccard similarity
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    async def _insert_screenshot_at_paragraph(
        self, 
        doc: Document, 
        screenshot: Screenshot, 
        paragraph_index: int, 
        task_id: int
    ):
        """Insert screenshot after the specified paragraph"""
        
        try:
            # Get the paragraph to insert after
            target_paragraph = doc.paragraphs[paragraph_index]
            
            # Add some spacing
            spacing_para = target_paragraph._element.getparent().insert_after(
                target_paragraph._element, 
                doc._body._element._new_paragraph()
            )
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_run = caption_para.add_run(f"Figure {task_id}: Output for Question {task_id}")
            caption_run.font.size = Inches(0.1)
            
            # Add the image
            img_para = doc.add_paragraph()
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Calculate appropriate image size
            max_width = Inches(6)  # Maximum width
            max_height = Inches(4)  # Maximum height
            
            # Scale image if needed
            width = min(max_width, Inches(screenshot.width / 100))
            height = min(max_height, Inches(screenshot.height / 100))
            
            # Maintain aspect ratio
            aspect_ratio = screenshot.width / screenshot.height
            if width / height > aspect_ratio:
                width = height * aspect_ratio
            else:
                height = width / aspect_ratio
            
            # Insert image
            img_para.add_run().add_picture(screenshot.file_path, width=width, height=height)
            
            # Add spacing after image
            doc.add_paragraph()
            
        except Exception as e:
            print(f"Error inserting screenshot: {str(e)}")
    
    async def get_report_info(self, file_path: str) -> dict:
        """Get information about a generated report"""
        try:
            if not os.path.exists(file_path):
                return {"exists": False}
            
            stat = os.stat(file_path)
            return {
                "exists": True,
                "file_size": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
            }
        except Exception as e:
            return {"exists": False, "error": str(e)}


# Global instance
docx_service = DocxService()
